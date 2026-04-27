import os
import io
import copy
import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

import config_loader
import user_manager

_TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'templates')
_TEMPLATE_FILE = os.path.join(_TEMPLATE_DIR, 'project_report_template.docx')

# Map status keys to human-readable labels
_STATUS_LABELS = {
    'done': 'Done',
    'not_ok': 'Not OK',
    'idle': 'Idle',
    'wip': 'Work In Progress',
}


def _get_project_data():
    """Collects all project information for the report."""
    master_projects = user_manager.get_master_projects()
    projects = []
    for mp in master_projects:
        assigned_to, status = user_manager.get_project_assignment(mp['name'])
        display_name = ''
        if assigned_to:
            display_name = user_manager.get_display_name(assigned_to)
        projects.append({
            'name': mp['name'],
            'folder': mp['folder'],
            'assigned_to': assigned_to or '',
            'display_name': display_name,
            'status': status or '',
            'status_label': _STATUS_LABELS.get(status, 'Unassigned'),
            'notes': mp.get('notes', ''),
        })
    return projects


def _clear_content_paragraphs(doc):
    """Removes content paragraphs (between cover and last page).
    Updated template structure (no TOC page):
      Paragraphs 0-35:  Cover page (section 0) — keep
      Paragraphs 36-132: Content               — CLEAR
      Paragraph  133:   Section break (section 1 boundary) — KEEP
      Paragraph  134:   Empty                              — keep
      Paragraph  135:   Section break (section 2 boundary) — KEEP
      Paragraph  136:   Last page with image               — keep
    """
    body = doc.element.body
    paras = doc.paragraphs

    # Remove content paragraphs 36 to 132 inclusive
    # Para 133 has the section break — must keep it
    elements_to_remove = []
    for i in range(36, min(133, len(paras))):
        elements_to_remove.append(paras[i]._element)

    for el in elements_to_remove:
        body.remove(el)


def _find_insert_point(doc):
    """Returns the XML element BEFORE which new content should be inserted.
    After clearing paras 36-132, the section-break paragraph (old 133) is now at index 36.
    We insert before it so content stays in section 1.
    """
    paras = doc.paragraphs
    if len(paras) > 36:
        return paras[36]._element
    return None


def _add_paragraph_before(doc, insert_before, text, style_name=None, bold=False, font_size=None, alignment=None, space_after=None):
    """Insert a new paragraph before a given element in the document body."""
    new_p = doc.add_paragraph()
    p_element = new_p._element
    # Move it before the insert point
    body = doc.element.body
    body.remove(p_element)
    if insert_before is not None:
        insert_before.addprevious(p_element)
    else:
        body.append(p_element)

    # Apply style if it exists
    if style_name:
        try:
            new_p.style = doc.styles[style_name]
        except KeyError:
            pass

    # Set text
    if text:
        run = new_p.add_run(text)
        if bold:
            run.bold = True
        if font_size:
            run.font.size = Pt(font_size)

    if alignment is not None:
        new_p.alignment = alignment

    if space_after is not None:
        new_p.paragraph_format.space_after = Pt(space_after)

    return new_p


def generate_project_report():
    """Generates the project report DOCX and returns it as bytes."""
    doc = Document(_TEMPLATE_FILE)
    now = datetime.datetime.now()
    date_str = now.strftime('%d/%m/%Y %H:%M')
    projects = _get_project_data()

    # ── 1. Cover page ──
    # Clear "A`" garbage text from paragraph 0
    if doc.paragraphs[0].runs:
        for run in doc.paragraphs[0].runs:
            if run.text.strip():
                run.text = ''

    # ── 2. Clear existing content paragraphs ──
    _clear_content_paragraphs(doc)

    # ── 3. Find insert point (before last page) ──
    insert_before = _find_insert_point(doc)

    # ── 4. Write clean project content: Header = "ProjectName — Status", body = notes ──
    for proj in projects:
        # Project heading: "AVINOR — Work In Progress"
        header_text = f'{proj["name"]} — {proj["status_label"]}'
        _add_paragraph_before(doc, insert_before, header_text,
                              style_name='00_Título Nivel 1')

        # Notes content
        notes = proj.get('notes', '').strip()
        if notes:
            lines = notes.split('\n')
            for line in lines:
                trimmed = line.strip()
                if not trimmed:
                    _add_paragraph_before(doc, insert_before, '',
                                          style_name='04_Cuerpo de texto')
                elif trimmed.startswith('-') or trimmed.startswith('•'):
                    # Bullet point
                    bullet_text = trimmed.lstrip('-•').strip()
                    _add_paragraph_before(doc, insert_before,
                                          f'• {bullet_text}',
                                          style_name='04_Cuerpo de texto')
                else:
                    _add_paragraph_before(doc, insert_before, trimmed,
                                          style_name='04_Cuerpo de texto')
        else:
            _add_paragraph_before(doc, insert_before, 'No notes available.',
                                  style_name='04_Cuerpo de texto')

        # Spacing between projects
        _add_paragraph_before(doc, insert_before, '',
                              style_name='04_Cuerpo de texto')

    # ── 5. Export to bytes ──
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue(), f'PULPITANS_Project_Report_{now.strftime("%Y%m%d_%H%M")}.docx'
